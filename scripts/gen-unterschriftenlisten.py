#!/usr/bin/env python3
"""
Generator für Zirkularbeschluss-Unterschriftenlisten pro STWEG.

Liest aktuelle Daten aus /tmp/units.json (vom DB-Export) und ersetzt die
Datenzeilen in den Template-DOCX-Dateien (/tmp/ul-template-stweg{N}.docx).
Erhält Header-Bereich, Tabellen-Styling und Stempel-Tabelle.

Nutzung:
    python3 gen-unterschriftenlisten.py [stwegN ...]

Erfordert:
    /tmp/units.json — JSON-Array von wohnungen mit eig/verw/adr-Aggregaten
    /tmp/ul-template-stweg{1..8}.docx — bestehende Templates

Schreibt: /tmp/ul-new-stweg{N}.docx mit aktualisierter Datentabelle + Footer.
"""
import docx, sys, os, json
from copy import deepcopy
from datetime import date

ALL_UNITS = json.load(open('/tmp/units.json'))


def fetch_units(stweg):
    rows = []
    for r in ALL_UNITS:
        if r['stweg'] != stweg:
            continue
        if stweg != 8 and r['typ'] not in ('Wohnung', 'Hobbyraum'):
            continue
        names_parts = []
        if r['eig']:
            names_parts.append(r['eig'])
        if r['verw']:
            names_parts.append(f"{r['verw']} (Verwalter)")
        eig_str = ', '.join(names_parts) if names_parts else ''
        wz, wn = r['wertquote_zaehler'], r['wertquote_nenner']
        wq_str = f"{wz}/{wn}" if (wz and wn) else '____'
        # Adresse: dedup nach Normalisierung
        addrs = [a.strip() for a in (r['adr'] or '').split(';') if a.strip()]
        seen, unique_addrs = set(), []
        for a in addrs:
            key = ' '.join(a.replace(',', '').split()).lower()
            if key not in seen:
                seen.add(key)
                unique_addrs.append(a)
        rows.append({
            'bezeichnung': r['bezeichnung'], 'typ': r['typ'],
            'eigentuemer': eig_str,
            'adresse': '; '.join(unique_addrs),
            'wertquote': wq_str,
            'wq_z': wz, 'wq_n': wn,
        })
    return rows


def find_data_table(doc):
    for t in doc.tables:
        if len(t.rows) > 1 and t.rows[0].cells[0].text.strip() == 'Einheit':
            return t
    return None


def set_footer(doc, stweg, total_units, datum_iso):
    """Footer: 'STWEG N — Unterschriftenliste · X Einheiten · Seite P von N · Stand {datum}'"""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from docx.shared import Pt, RGBColor
    for section in doc.sections:
        footer = section.footer
        for p in list(footer.paragraphs):
            elem = p._element
            elem.getparent().remove(elem)
        para = footer.add_paragraph()
        para.alignment = 1  # center
        para.add_run(f'STWEG {stweg} — Unterschriftenliste · {total_units} Einheiten · Seite ')
        # PAGE field
        r = para.add_run()
        fld1 = OxmlElement('w:fldChar'); fld1.set(qn('w:fldCharType'), 'begin')
        instr = OxmlElement('w:instrText'); instr.text = ' PAGE '
        fld2 = OxmlElement('w:fldChar'); fld2.set(qn('w:fldCharType'), 'end')
        r._r.append(fld1); r._r.append(instr); r._r.append(fld2)
        para.add_run(' von ')
        # NUMPAGES field
        r2 = para.add_run()
        fld3 = OxmlElement('w:fldChar'); fld3.set(qn('w:fldCharType'), 'begin')
        instr2 = OxmlElement('w:instrText'); instr2.text = ' NUMPAGES '
        fld4 = OxmlElement('w:fldChar'); fld4.set(qn('w:fldCharType'), 'end')
        r2._r.append(fld3); r2._r.append(instr2); r2._r.append(fld4)
        para.add_run(f' · Stand {datum_iso}')
        for run in para.runs:
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)


def set_cell_text_preserving_style(cell, text):
    if not cell.paragraphs:
        cell.text = text
        return
    para = cell.paragraphs[0]
    if not para.runs:
        para.text = text
        return
    first_run = para.runs[0]
    first_run.text = text
    for r in para.runs[1:]:
        r.text = ''
    for p in cell.paragraphs[1:]:
        p_elem = p._element
        p_elem.getparent().remove(p_elem)


def update_unterschriftenliste(template_path, output_path, stweg, units):
    doc = docx.Document(template_path)
    table = find_data_table(doc)
    if table is None:
        raise RuntimeError(f"Daten-Tabelle (Header 'Einheit') nicht gefunden in {template_path}")
    tbl = table._tbl
    if len(table.rows) < 3:
        raise RuntimeError("Tabelle zu klein — erwartet mindestens header + 1 data + total")
    template_row = deepcopy(table.rows[1]._tr)
    total_row = deepcopy(table.rows[-1]._tr)
    rows_to_remove = list(tbl.iterchildren('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tr'))[1:]
    for r in rows_to_remove:
        tbl.remove(r)
    sum_wq = 0
    for u in units:
        new_tr = deepcopy(template_row)
        tbl.append(new_tr)
        new_row = table.rows[-1]
        cells = new_row.cells
        set_cell_text_preserving_style(cells[0], u['bezeichnung'])
        set_cell_text_preserving_style(cells[1], u['eigentuemer'])
        set_cell_text_preserving_style(cells[2], u['adresse'])
        set_cell_text_preserving_style(cells[3], u['wertquote'])
        set_cell_text_preserving_style(cells[4], '☐ JA    ☐ NEIN')
        set_cell_text_preserving_style(cells[5], '')
        if u.get('wq_z'):
            sum_wq += u['wq_z']
    tbl.append(total_row)
    new_total_row = table.rows[-1]
    set_cell_text_preserving_style(new_total_row.cells[0], 'TOTAL')
    set_cell_text_preserving_style(new_total_row.cells[1], f"{len(units)} Einheiten")
    set_cell_text_preserving_style(new_total_row.cells[2], '')
    if stweg == 8:
        set_cell_text_preserving_style(new_total_row.cells[3], '')
    else:
        set_cell_text_preserving_style(new_total_row.cells[3], f"{sum_wq}/1000")
    set_cell_text_preserving_style(new_total_row.cells[4], '')
    set_cell_text_preserving_style(new_total_row.cells[5], '')
    # Footer mit Seitenzahlen
    set_footer(doc, stweg, len(units), date.today().isoformat())
    doc.save(output_path)
    return sum_wq


def main():
    stwegs = [int(s) for s in sys.argv[1:]] if len(sys.argv) > 1 else list(range(1, 9))
    for s in stwegs:
        units = fetch_units(s)
        tpl = f"/tmp/ul-template-stweg{s}.docx"
        out = f"/tmp/ul-new-stweg{s}.docx"
        sum_wq = update_unterschriftenliste(tpl, out, s, units)
        print(f"STWEG {s}: {len(units)} units, Σ wertquoten={sum_wq}/1000 → {out}")


if __name__ == '__main__':
    main()
